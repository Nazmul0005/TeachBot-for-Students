"""
Enhanced Document Processing for TeachBot RAG System
====================================================

This script creates vector embeddings from multiple document formats:
- PDF files: Using PyMuPDF for better Bengali/multilingual text extraction
- DOCX files: Using python-docx for comprehensive content extraction (paragraphs + tables)
- TXT files: With multiple encoding support (UTF-8, Latin-1, etc.)

Features:
- Automatic file type detection in data folder
- Improved text extraction for Bengali/Bangla content
- Fallback methods for robust processing
- Comprehensive error handling
- Detailed progress reporting

Supported file formats: .pdf, .docx, .txt
Output: FAISS vector store with OpenAI embeddings
"""

import os
import glob
from langchain_community.document_loaders import PyPDFLoader, DirectoryLoader, Docx2txtLoader, TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
import fitz  # PyMuPDF for better text extraction
from langchain.schema import Document
from docx import Document as DocxDocument

## Uncomment the following files if you're not using pipenv as your virtual environment manager
from dotenv import load_dotenv, find_dotenv
load_dotenv(find_dotenv())

# Get OpenAI API key
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")

# Step 1: Load multiple file formats (PDF, DOCX, TXT) with better text extraction
DATA_PATH = "data/"

def load_pdf_files_improved(pdf_files):
    """Load PDFs with better text extraction for Bangla text"""
    documents = []
    
    for pdf_file in pdf_files:
        print(f"Processing PDF: {pdf_file}")
        
        try:
            # Use PyMuPDF for better text extraction
            doc = fitz.open(pdf_file)
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # Extract text with better encoding handling
                text = page.get_text("text", flags=fitz.TEXT_PRESERVE_WHITESPACE)
                
                # Clean up the text
                text = text.strip()
                
                if text:  # Only add non-empty pages
                    # Create document with metadata
                    metadata = {
                        'source': pdf_file,
                        'file_type': 'pdf',
                        'page': page_num,
                        'page_label': str(page_num + 1),
                        'total_pages': len(doc)
                    }
                    
                    document = Document(page_content=text, metadata=metadata)
                    documents.append(document)
                    
                    # Print first few characters to check encoding
                    print(f"  Page {page_num + 1}: {text[:100]}...")
            
            doc.close()
            print(f"  ✓ Successfully processed {len(doc)} pages from {pdf_file}")
            
        except Exception as e:
            print(f"  ❌ Error processing {pdf_file}: {e}")
    
    return documents

def load_docx_files(docx_files):
    """Load DOCX files with better text extraction for Bangla text"""
    documents = []
    
    for docx_file in docx_files:
        print(f"Processing DOCX: {docx_file}")
        
        try:
            # Method 1: Using python-docx for better control
            doc = DocxDocument(docx_file)
            full_text = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text)
            
            text = '\n'.join(full_text).strip()
            
            if text:
                metadata = {
                    'source': docx_file,
                    'file_type': 'docx',
                    'page': 0,
                    'page_label': 'Document',
                    'total_pages': 1
                }
                
                document = Document(page_content=text, metadata=metadata)
                documents.append(document)
                
                print(f"  ✓ Successfully processed {docx_file}")
                print(f"  Content preview: {text[:100]}...")
            else:
                print(f"  ⚠ No text content found in {docx_file}")
                
        except Exception as e:
            print(f"  ❌ Error processing {docx_file}: {e}")
            # Fallback to langchain loader
            try:
                print(f"  Trying fallback method for {docx_file}")
                loader = Docx2txtLoader(docx_file)
                docs = loader.load()
                documents.extend(docs)
                print(f"  ✓ Fallback method succeeded for {docx_file}")
            except Exception as e2:
                print(f"  ❌ Fallback method also failed: {e2}")
    
    return documents

def load_txt_files(txt_files):
    """Load TXT files with proper encoding handling"""
    documents = []
    
    for txt_file in txt_files:
        print(f"Processing TXT: {txt_file}")
        
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
            text = None
            
            for encoding in encodings:
                try:
                    with open(txt_file, 'r', encoding=encoding) as f:
                        text = f.read().strip()
                    print(f"  ✓ Successfully read with {encoding} encoding")
                    break
                except UnicodeDecodeError:
                    continue
            
            if text:
                metadata = {
                    'source': txt_file,
                    'file_type': 'txt',
                    'page': 0,
                    'page_label': 'Document',
                    'total_pages': 1
                }
                
                document = Document(page_content=text, metadata=metadata)
                documents.append(document)
                
                print(f"  ✓ Successfully processed {txt_file}")
                print(f"  Content preview: {text[:100]}...")
            else:
                print(f"  ❌ Could not read {txt_file} with any encoding")
                
        except Exception as e:
            print(f"  ❌ Error processing {txt_file}: {e}")
    
    return documents

def load_all_documents(data_path):
    """Load all supported document types from the data directory"""
    print(f"🔍 Scanning directory: {data_path}")
    
    # Find all supported files
    pdf_files = glob.glob(os.path.join(data_path, "*.pdf"))
    docx_files = glob.glob(os.path.join(data_path, "*.docx"))
    txt_files = glob.glob(os.path.join(data_path, "*.txt"))
    
    print(f"Found files:")
    print(f"  📄 PDF files: {len(pdf_files)}")
    print(f"  📝 DOCX files: {len(docx_files)}")
    print(f"  📃 TXT files: {len(txt_files)}")
    
    all_documents = []
    
    # Load PDF files
    if pdf_files:
        print(f"\n📄 Processing {len(pdf_files)} PDF files...")
        pdf_docs = load_pdf_files_improved(pdf_files)
        all_documents.extend(pdf_docs)
        print(f"✓ Loaded {len(pdf_docs)} documents from PDF files")
    
    # Load DOCX files
    if docx_files:
        print(f"\n📝 Processing {len(docx_files)} DOCX files...")
        docx_docs = load_docx_files(docx_files)
        all_documents.extend(docx_docs)
        print(f"✓ Loaded {len(docx_docs)} documents from DOCX files")
    
    # Load TXT files
    if txt_files:
        print(f"\n📃 Processing {len(txt_files)} TXT files...")
        txt_docs = load_txt_files(txt_files)
        all_documents.extend(txt_docs)
        print(f"✓ Loaded {len(txt_docs)} documents from TXT files")
    
    return all_documents

def load_documents_fallback(data_path):
    """Fallback method using DirectoryLoader for all file types"""
    print("Using fallback DirectoryLoader method...")
    documents = []
    
    # Load PDFs
    try:
        pdf_loader = DirectoryLoader(data_path, glob='*.pdf', loader_cls=PyPDFLoader)
        pdf_docs = pdf_loader.load()
        documents.extend(pdf_docs)
        print(f"✓ Loaded {len(pdf_docs)} PDF documents via fallback")
    except Exception as e:
        print(f"⚠ PDF fallback failed: {e}")
    
    # Load DOCX files
    try:
        docx_loader = DirectoryLoader(data_path, glob='*.docx', loader_cls=Docx2txtLoader)
        docx_docs = docx_loader.load()
        documents.extend(docx_docs)
        print(f"✓ Loaded {len(docx_docs)} DOCX documents via fallback")
    except Exception as e:
        print(f"⚠ DOCX fallback failed: {e}")
    
    # Load TXT files
    try:
        txt_loader = DirectoryLoader(data_path, glob='*.txt', loader_cls=TextLoader)
        txt_docs = txt_loader.load()
        documents.extend(txt_docs)
        print(f"✓ Loaded {len(txt_docs)} TXT documents via fallback")
    except Exception as e:
        print(f"⚠ TXT fallback failed: {e}")
    
    return documents

# Load all documents with improved methods, fallback if needed
try:
    documents = load_all_documents(DATA_PATH)
    print(f"\n✓ Successfully loaded {len(documents)} documents using improved methods")
except Exception as e:
    print(f"\n⚠ Improved method failed: {e}")
    print("Using fallback method...")
    documents = load_documents_fallback(DATA_PATH)
    print(f"✓ Loaded {len(documents)} documents using fallback methods")

print(f"\n📊 Total documents loaded: {len(documents)}")

# Print sample content to verify text extraction
if documents:
    print("\n--- Sample extracted text ---")
    print(documents[0].page_content[:500])
    print("--- End sample ---\n")

# Step 2: Create Chunks (optimized for Bangla)
def create_chunks(extracted_data):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1500,  # Larger chunks for better context
        chunk_overlap=200,  # More overlap for continuity
        separators=["\n\n", "\n", "।", ".", " ", ""]  # Bangla-aware separators
    )
    text_chunks = text_splitter.split_documents(extracted_data)
    return text_chunks

text_chunks = create_chunks(extracted_data=documents)
print("Length of Text Chunks: ", len(text_chunks))

# Print sample chunks to verify
if text_chunks:
    print("\n--- Sample chunk ---")
    print(text_chunks[0].page_content)
    print("--- End sample chunk ---\n")

# Step 3: Create Vector Embeddings using OpenAI (better for multilingual)
def get_embedding_model():
    embedding_model = OpenAIEmbeddings(
        model="text-embedding-3-large",  # Latest and best OpenAI embedding model
        openai_api_key=OPENAI_API_KEY
    )
    return embedding_model

embedding_model = get_embedding_model()

# Step 4: Store embeddings in FAISS
DB_FAISS_PATH = "vectorstore/db_faiss_openai_improved"  # New path for improved version
db = FAISS.from_documents(text_chunks, embedding_model)
db.save_local(DB_FAISS_PATH)
print(f"✓ Vector database saved to {DB_FAISS_PATH}")

# Step 5: Test the vector store
print("\n--- Testing vector store ---")
test_query = "রবীন্দ্রনাথ ঠাকুর"
retriever = db.as_retriever(search_kwargs={'k': 3})
docs = retriever.get_relevant_documents(test_query)

print(f"Found {len(docs)} relevant documents for '{test_query}':")
for i, doc in enumerate(docs):
    print(f"\nDocument {i+1}:")
    print(f"Content: {doc.page_content[:200]}...")
    print(f"Source: {doc.metadata.get('source', 'Unknown')}, Page: {doc.metadata.get('page', 'Unknown')}")